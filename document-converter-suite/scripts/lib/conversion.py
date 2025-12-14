from __future__ import annotations

import shutil
from pathlib import Path
from typing import List, Optional

from docx import Document as DocxDocument

from .docx_reader import read_docx_content
from .docx_writer import write_docx_from_text
from .pdf_reader import read_pdf_text
from .pdf_writer import write_pdf_from_sections
from .pptx_reader import read_pptx_content
from .pptx_writer import write_pptx_from_slides
from .types import SUPPORTED_INPUT_EXTS, SUPPORTED_OUTPUT_EXTS
from .utils import chunk_list, guess_bullets, safe_cell_value
from .xlsx_reader import read_xlsx_content
from .xlsx_writer import write_xlsx_from_sheets


def convert_document(
    input_path: Path,
    output_path: Path,
    max_pages: int = 200,
    max_chars: int = 300000,
    max_rows: int = 200,
    max_cols: int = 50,
    verbose: bool = False,
) -> None:
    """Convert a single document between PDF, DOCX, PPTX, XLSX.

    This is best-effort conversion focused on text and basic tables.
    """
    in_ext = input_path.suffix.lower()
    out_ext = output_path.suffix.lower()

    if in_ext not in SUPPORTED_INPUT_EXTS:
        raise ValueError(f"Unsupported input extension: {in_ext}")
    if out_ext not in SUPPORTED_OUTPUT_EXTS:
        raise ValueError(f"Unsupported output extension: {out_ext}")

    # Shortcut: same extension -> copy
    if in_ext == out_ext:
        if verbose:
            print(f"Copying (same format): {input_path} -> {output_path}")
        shutil.copy2(str(input_path), str(output_path))
        return

    # === Convert to DOCX ===
    if out_ext == ".docx":
        if in_ext == ".pdf":
            _pdf_to_docx(input_path, output_path, max_pages=max_pages, max_chars=max_chars)
            return
        if in_ext == ".pptx":
            _pptx_to_docx(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".xlsx":
            _xlsx_to_docx(input_path, output_path, max_rows=max_rows, max_cols=max_cols)
            return

    # === Convert to PPTX ===
    if out_ext == ".pptx":
        if in_ext == ".pdf":
            _pdf_to_pptx(input_path, output_path, max_pages=max_pages, max_chars=max_chars)
            return
        if in_ext == ".docx":
            _docx_to_pptx(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".xlsx":
            _xlsx_to_pptx(input_path, output_path, max_rows=max_rows, max_cols=max_cols)
            return

    # === Convert to XLSX ===
    if out_ext == ".xlsx":
        if in_ext == ".pdf":
            _pdf_to_xlsx(input_path, output_path, max_pages=max_pages, max_chars=max_chars)
            return
        if in_ext == ".docx":
            _docx_to_xlsx(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".pptx":
            _pptx_to_xlsx(input_path, output_path, max_chars=max_chars)
            return

    # === Convert to PDF ===
    if out_ext == ".pdf":
        if in_ext == ".docx":
            _docx_to_pdf(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".pptx":
            _pptx_to_pdf(input_path, output_path, max_chars=max_chars)
            return
        if in_ext == ".xlsx":
            _xlsx_to_pdf(input_path, output_path, max_rows=max_rows, max_cols=max_cols)
            return

    raise ValueError(f"No conversion path implemented: {in_ext} -> {out_ext}")


# --------------------------
# PDF -> *
# --------------------------

def _pdf_to_docx(input_path: Path, output_path: Path, max_pages: int, max_chars: int) -> None:
    pdf = read_pdf_text(input_path, max_pages=max_pages, max_chars=max_chars)
    doc = DocxDocument()
    doc.add_heading(input_path.stem, level=1)
    for i, page_text in enumerate(pdf.pages, start=1):
        doc.add_heading(f"Page {i}", level=2)
        for ln in page_text.splitlines():
            ln = (ln or "").strip()
            if not ln:
                continue
            doc.add_paragraph(ln)
        if i < len(pdf.pages):
            doc.add_page_break()
    doc.save(str(output_path))


def _pdf_to_pptx(input_path: Path, output_path: Path, max_pages: int, max_chars: int) -> None:
    pdf = read_pdf_text(input_path, max_pages=max_pages, max_chars=max_chars)
    slides: List[dict] = []
    for i, page_text in enumerate(pdf.pages, start=1):
        slides.append({"title": f"Page {i}", "bullets": guess_bullets(page_text)})
    write_pptx_from_slides(output_path, slides, title=input_path.stem)


def _pdf_to_xlsx(input_path: Path, output_path: Path, max_pages: int, max_chars: int) -> None:
    pdf = read_pdf_text(input_path, max_pages=max_pages, max_chars=max_chars)
    rows: List[List[str]] = []
    for i, page_text in enumerate(pdf.pages, start=1):
        rows.append([f"=== Page {i} ==="])
        for ln in page_text.splitlines():
            ln = (ln or "").strip()
            if ln:
                rows.append([ln])
        rows.append([""])
    write_xlsx_from_sheets(output_path, [{"name": "PDF Text", "rows": rows}])


# --------------------------
# DOCX -> *
# --------------------------

def _docx_to_pptx(input_path: Path, output_path: Path, max_chars: int) -> None:
    content = read_docx_content(input_path, max_chars=max_chars)

    # Heuristic: if there are headings, start slides at headings.
    slides: List[dict] = []
    current_title: Optional[str] = None
    current_bullets: List[str] = []

    def flush():
        nonlocal current_title, current_bullets
        if current_title or current_bullets:
            slides.append({"title": current_title or "", "bullets": current_bullets})
        current_title = None
        current_bullets = []

    for p in content.paragraphs:
        # Treat lines that match known headings as slide boundaries
        if p in set(content.headings):
            flush()
            current_title = p
            continue
        # Chunk big docs into manageable slides
        current_bullets.append(p)
        if len(current_bullets) >= 10:
            flush()

    flush()
    if not slides:
        slides = [{"title": input_path.stem, "bullets": content.paragraphs[:10]}]

    # Normalize bullets a bit
    for s in slides:
        s["bullets"] = [b for b in guess_bullets("\n".join(s.get("bullets") or []), max_lines=12)]

    write_pptx_from_slides(output_path, slides, title=input_path.stem)


def _docx_to_xlsx(input_path: Path, output_path: Path, max_chars: int) -> None:
    content = read_docx_content(input_path, max_chars=max_chars)

    sheets: List[dict] = []

    # Text sheet
    text_rows = [[p] for p in content.paragraphs]
    sheets.append({"name": "Text", "rows": text_rows})

    # Tables -> their own sheets
    for idx, t in enumerate(content.tables, start=1):
        sheets.append({"name": f"Table{idx}", "rows": t.rows})

    write_xlsx_from_sheets(output_path, sheets)


def _docx_to_pdf(input_path: Path, output_path: Path, max_chars: int) -> None:
    content = read_docx_content(input_path, max_chars=max_chars)
    tables = [t.rows for t in content.tables]
    sections = [{"heading": input_path.stem, "paragraphs": content.paragraphs, "tables": tables}]
    write_pdf_from_sections(output_path, title=input_path.stem, sections=sections)


# --------------------------
# PPTX -> *
# --------------------------

def _pptx_to_docx(input_path: Path, output_path: Path, max_chars: int) -> None:
    content = read_pptx_content(input_path, max_chars=max_chars)

    doc = DocxDocument()
    doc.add_heading(input_path.stem, level=1)

    for s in content.slides:
        heading = s.title or f"Slide {s.index}"
        doc.add_heading(heading, level=2)

        # Text blocks -> bullets-ish
        for block in s.texts:
            for bullet in guess_bullets(block, max_lines=20):
                doc.add_paragraph(bullet, style="List Bullet")

        # Tables
        for t_i, t in enumerate(s.tables, start=1):
            if not t:
                continue
            doc.add_paragraph("")
            doc.add_heading(f"Table {t_i}", level=3)
            n_rows = len(t)
            n_cols = max((len(r) for r in t), default=0)
            if n_cols <= 0:
                continue
            table = doc.add_table(rows=n_rows, cols=n_cols)
            for r_i, row in enumerate(t):
                for c_i in range(n_cols):
                    table.cell(r_i, c_i).text = safe_cell_value(row[c_i] if c_i < len(row) else "")

    doc.save(str(output_path))


def _pptx_to_xlsx(input_path: Path, output_path: Path, max_chars: int) -> None:
    content = read_pptx_content(input_path, max_chars=max_chars)

    # Slides sheet: one row per slide with combined text
    rows: List[List[str]] = [["Slide", "Title", "Text"]]
    for s in content.slides:
        combined = "\n\n".join(s.texts)
        rows.append([str(s.index), s.title, combined])

    sheets: List[dict] = [{"name": "Slides", "rows": rows}]

    # Any tables: each table becomes its own sheet (bounded name)
    t_counter = 1
    for s in content.slides:
        for t in s.tables:
            name = f"Table{t_counter}"[:31]
            sheets.append({"name": name, "rows": t})
            t_counter += 1

    write_xlsx_from_sheets(output_path, sheets)


def _pptx_to_pdf(input_path: Path, output_path: Path, max_chars: int) -> None:
    content = read_pptx_content(input_path, max_chars=max_chars)
    sections: List[dict] = []
    for s in content.slides:
        heading = s.title or f"Slide {s.index}"
        paras: List[str] = []
        for block in s.texts:
            paras.extend(guess_bullets(block, max_lines=25))
        tables = s.tables
        sections.append({"heading": heading, "paragraphs": paras, "tables": tables})
    write_pdf_from_sections(output_path, title=input_path.stem, sections=sections)


# --------------------------
# XLSX -> *
# --------------------------

def _xlsx_to_docx(input_path: Path, output_path: Path, max_rows: int, max_cols: int) -> None:
    content = read_xlsx_content(input_path, max_rows=max_rows, max_cols=max_cols)

    doc = DocxDocument()
    doc.add_heading(input_path.stem, level=1)

    for sh in content.sheets:
        doc.add_heading(sh.name, level=2)
        if not sh.cells:
            doc.add_paragraph("(empty sheet)")
            continue
        rows = sh.cells
        n_rows = len(rows)
        n_cols = max((len(r) for r in rows), default=0)
        table = doc.add_table(rows=n_rows, cols=n_cols)
        for r_i, row in enumerate(rows):
            for c_i in range(n_cols):
                table.cell(r_i, c_i).text = safe_cell_value(row[c_i] if c_i < len(row) else "")
        doc.add_paragraph("")

    doc.save(str(output_path))


def _xlsx_to_pptx(input_path: Path, output_path: Path, max_rows: int, max_cols: int) -> None:
    content = read_xlsx_content(input_path, max_rows=max_rows, max_cols=max_cols)

    slides: List[dict] = []
    for sh in content.sheets:
        rows = sh.cells
        # If small enough, push it as a table; otherwise summarize first column as bullets.
        if len(rows) <= 20 and (max((len(r) for r in rows), default=0) <= 10):
            slides.append({"title": sh.name, "bullets": [], "tables": [rows]})
        else:
            first_col = [r[0] for r in rows if r and r[0]]
            bullets = first_col[:12] if first_col else ["(sheet too large to render as table)"]
            slides.append({"title": sh.name, "bullets": bullets})

    write_pptx_from_slides(output_path, slides, title=input_path.stem)


def _xlsx_to_pdf(input_path: Path, output_path: Path, max_rows: int, max_cols: int) -> None:
    content = read_xlsx_content(input_path, max_rows=max_rows, max_cols=max_cols)
    sections: List[dict] = []
    for sh in content.sheets:
        sections.append({"heading": sh.name, "paragraphs": [], "tables": [sh.cells]})
    write_pdf_from_sections(output_path, title=input_path.stem, sections=sections)

