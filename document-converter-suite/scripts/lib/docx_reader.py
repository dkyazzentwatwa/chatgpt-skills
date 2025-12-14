from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import List

from docx import Document


@dataclass
class DocxTable:
    rows: List[List[str]]


@dataclass
class DocxContent:
    paragraphs: List[str]
    headings: List[str]
    tables: List[DocxTable]


def read_docx_content(path: Path, max_chars: int = 300000) -> DocxContent:
    """Read text and basic tables from a .docx.

    Heuristics:
      - Heading detection: paragraph style name starts with 'Heading'
      - Tables: raw cell text
    """
    doc = Document(str(path))
    paras: List[str] = []
    headings: List[str] = []

    budget = max_chars
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt:
            continue
        try:
            style_name = p.style.name if p.style else ""
        except Exception:
            style_name = ""
        if style_name.lower().startswith("heading"):
            headings.append(txt)
        if budget <= 0:
            break
        if len(txt) > budget:
            txt = txt[:budget]
        budget -= len(txt)
        paras.append(txt)

    tables: List[DocxTable] = []
    for t in doc.tables:
        rows: List[List[str]] = []
        for row in t.rows:
            rows.append([(cell.text or "").strip() for cell in row.cells])
        if rows:
            tables.append(DocxTable(rows=rows))

    return DocxContent(paragraphs=paras, headings=headings, tables=tables)
