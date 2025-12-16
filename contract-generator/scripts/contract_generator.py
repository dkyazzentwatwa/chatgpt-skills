#!/usr/bin/env python3
"""Contract Generator - Generate legal contracts from templates."""

import argparse
import json
import os
import re
from pathlib import Path
from typing import Dict
from docx import Document
import pandas as pd


class ContractGenerator:
    """Generate contracts from DOCX templates."""

    def __init__(self):
        self.template = None
        self.variables = {}

    def load_template(self, filepath: str):
        """Load DOCX template."""
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Template not found: {filepath}")
        self.template = Document(filepath)
        return self

    def set_variables(self, variables: Dict[str, str]):
        """Set template variables."""
        self.variables.update(variables)
        return self

    def _replace_placeholders(self, text: str) -> str:
        """Replace {{variable}} placeholders with values."""
        for key, value in self.variables.items():
            placeholder = f"{{{{{key}}}}}"
            text = text.replace(placeholder, str(value))
        return text

    def generate(self):
        """Generate contract by replacing placeholders."""
        if self.template is None:
            raise ValueError("No template loaded")

        # Replace in paragraphs
        for paragraph in self.template.paragraphs:
            for run in paragraph.runs:
                run.text = self._replace_placeholders(run.text)

        # Replace in tables
        for table in self.template.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.text = self._replace_placeholders(run.text)

        return self

    def validate(self) -> list:
        """Check for unreplaced placeholders."""
        missing = []
        pattern = r'\{\{(\w+)\}\}'

        for paragraph in self.template.paragraphs:
            matches = re.findall(pattern, paragraph.text)
            missing.extend(matches)

        for table in self.template.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(pattern, paragraph.text)
                        missing.extend(matches)

        return list(set(missing))

    def save(self, output: str):
        """Save generated contract."""
        if self.template is None:
            raise ValueError("No template loaded")

        os.makedirs(os.path.dirname(output) or '.', exist_ok=True)
        self.template.save(output)
        return output

    def batch_generate(self, csv_path: str, output_dir: str, template_path: str):
        """Generate multiple contracts from CSV."""
        df = pd.read_csv(csv_path)
        os.makedirs(output_dir, exist_ok=True)

        outputs = []
        for i, row in df.iterrows():
            print(f"Generating contract {i+1}/{len(df)}...")

            # Load fresh template
            self.load_template(template_path)

            # Set variables from row
            self.set_variables(row.to_dict())

            # Generate
            self.generate()

            # Save
            filename = f"contract_{i+1:04d}.docx"
            if 'output_filename' in row:
                filename = row['output_filename']

            output_path = os.path.join(output_dir, filename)
            self.save(output_path)
            outputs.append(output_path)

        return outputs


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Generate contracts from templates')
    parser.add_argument('--template', required=True, help='Template DOCX file')
    parser.add_argument('--vars', help='JSON file with variables')
    parser.add_argument('--csv', help='CSV file for batch generation')
    parser.add_argument('--output', help='Output file (single mode)')
    parser.add_argument('--output-dir', help='Output directory (batch mode)')

    args = parser.parse_args()

    generator = ContractGenerator()

    if args.csv:
        # Batch mode
        if not args.output_dir:
            print("Error: --output-dir required for batch mode")
            exit(1)

        outputs = generator.batch_generate(
            csv_path=args.csv,
            output_dir=args.output_dir,
            template_path=args.template
        )
        print(f"\n✓ Generated {len(outputs)} contracts in {args.output_dir}")

    else:
        # Single mode
        if not args.output:
            print("Error: --output required for single mode")
            exit(1)

        generator.load_template(args.template)

        if args.vars:
            with open(args.vars) as f:
                variables = json.load(f)
            generator.set_variables(variables)

        generator.generate()

        # Validate
        missing = generator.validate()
        if missing:
            print(f"Warning: Unreplaced placeholders: {missing}")

        generator.save(args.output)
        print(f"✓ Contract generated: {args.output}")
